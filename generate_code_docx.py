# generate_code_docx.py
# Script untuk membuat dokumen DOCX berisi kode Fuzzy AHP, Rule-Based, dan IndoBERT

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime
import os

# ─── Paths ───
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT_PATH = os.path.join(BASE_DIR, f"Kode_FuzzyAHP_RuleBased_IndoBERT_{datetime.now().strftime('%Y-%m-%dT%H-%M-%S')}.docx")

# Source code files
FILES = [
    {
        "section": "1. Pembobotan Fuzzy AHP",
        "file": os.path.join(BASE_DIR, "src", "lib", "dss", "fuzzyAHP.ts"),
        "lang": "TypeScript",
        "description": (
            "Modul ini mengimplementasikan metode Fuzzy Analytical Hierarchy Process (Fuzzy AHP) "
            "untuk pembobotan kriteria pada sistem penilaian risiko halal. "
            "Fuzzy AHP menggunakan Triangular Fuzzy Number (TFN) untuk menangani ketidakpastian "
            "dalam penilaian perbandingan berpasangan antar kriteria.\n\n"
            "Fungsi utama meliputi:\n"
            "• Perhitungan Fuzzy Synthetic Extent (FSE)\n"
            "• Defuzzifikasi menggunakan Center of Area (CoA)\n"
            "• Normalisasi bobot\n"
            "• Perhitungan Consistency Ratio (CR) berdasarkan tabel RI Saaty\n"
            "• Integrasi dengan database PostgreSQL via Prisma ORM untuk menyimpan dan membaca matriks perbandingan berpasangan"
        ),
    },
    {
        "section": "2. Rule-Based Risk Assessment Engine",
        "file": os.path.join(BASE_DIR, "src", "lib", "dss", "rule-engine.ts"),
        "lang": "TypeScript",
        "description": (
            "Modul Rule-Based Engine ini mengevaluasi tingkat risiko halal menggunakan pendekatan "
            "berbasis aturan (rule-based) dengan data dari file JSON Rule Base.\n\n"
            "Sistem menggunakan agregasi weakest-link (MAX):\n"
            "• Indikator → Konstruk: MAX dari semua skor indikator (I1..I5)\n"
            "• Konstruk → Stage: MAX dari semua konstruk dalam satu stage (CP)\n"
            "• Stage → Overall: MAX dari semua stage (CP1..CP9)\n\n"
            "Setiap indikator dinilai pada skala 1-5 (Sangat Rendah - Sangat Tinggi) "
            "dengan performance descriptor dan recommended action yang spesifik."
        ),
    },
    {
        "section": "3. IndoBERT Intent Classifier",
        "subsections": [
            {
                "subtitle": "3.1. Inference (Runtime) — intent-classifier.ts",
                "file": os.path.join(BASE_DIR, "src", "lib", "ml", "intent-classifier.ts"),
                "lang": "TypeScript",
                "description": (
                    "Modul ini merupakan kode inference untuk mengklasifikasikan intent pengguna "
                    "menggunakan model IndoBERT yang sudah di-fine-tune. Model di-load dari Hugging Face Hub "
                    "(NurfauzanDaffa/indobert-intent) menggunakan library @huggingface/transformers.\n\n"
                    "Fitur utama:\n"
                    "• Singleton pattern untuk efisiensi memori (model hanya di-load sekali)\n"
                    "• Fallback handling jika model gagal di-load\n"
                    "• Support Vercel deployment dengan konfigurasi cache khusus"
                ),
            },
            {
                "subtitle": "3.2. Training (Fine-tuning) — train_indobert_onnx_v2.ipynb",
                "file": os.path.join(BASE_DIR, "ml", "train_indobert_onnx_v2.ipynb"),
                "lang": "Python (Jupyter Notebook)",
                "description": (
                    "Notebook ini berisi kode training/fine-tuning model IndoBERT untuk klasifikasi intent. "
                    "Model dasar yang digunakan adalah indobenchmark/indobert-base-p1 yang di-fine-tune "
                    "menggunakan dataset intent khusus.\n\n"
                    "Langkah-langkah:\n"
                    "1. Install dependencies (transformers, datasets, optimum)\n"
                    "2. Login ke Hugging Face Hub\n"
                    "3. Load dataset CSV (train/test split)\n"
                    "4. Tokenisasi menggunakan IndoBERT tokenizer (max_length=128)\n"
                    "5. Fine-tuning dengan hyperparameter: lr=2e-5, batch=8, epochs=3\n"
                    "6. Ekspor model ke format ONNX\n"
                    "7. Upload ke Hugging Face Hub\n\n"
                    "Label intent yang dikenali:\n"
                    "• batch_trace, greeting, knowledge_query, operational_data, out_of_scope, risk_check"
                ),
            },
        ],
    },
    {
        "section": "4. Unit Test — Fuzzy AHP",
        "file": os.path.join(BASE_DIR, "src", "lib", "dss", "fuzzyAHP.test.ts"),
        "lang": "TypeScript",
        "description": (
            "File test ini berisi unit test untuk memvalidasi fungsi-fungsi inti Fuzzy AHP "
            "menggunakan framework Vitest. Test mencakup:\n"
            "• Perhitungan resiprokal TFN\n"
            "• Penjumlahan TFN\n"
            "• Defuzzifikasi Center of Area\n"
            "• Normalisasi bobot\n"
            "• Penentuan risk level\n"
            "• Validasi Consistency Ratio pada matriks yang perfectly consistent"
        ),
    },
]


def extract_notebook_code(filepath):
    """Extract Python code cells from Jupyter Notebook (.ipynb)."""
    import json
    with open(filepath, 'r', encoding='utf-8') as f:
        nb = json.load(f)
    
    code_blocks = []
    for cell in nb.get('cells', []):
        if cell['cell_type'] == 'code':
            source = ''.join(cell['source'])
            if source.strip():
                code_blocks.append(source)
        elif cell['cell_type'] == 'markdown':
            source = ''.join(cell['source'])
            if source.strip():
                code_blocks.append(f"# [Markdown Cell]\n# {source.replace(chr(10), chr(10) + '# ')}")
    
    return '\n\n'.join(code_blocks)


def create_code_style(doc):
    """Create a monospaced code style for the document."""
    style = doc.styles.add_style('CodeBlock', WD_STYLE_TYPE.PARAGRAPH)
    style.font.name = 'Consolas'
    style.font.size = Pt(8.5)
    style.font.color.rgb = RGBColor(30, 30, 30)
    style.paragraph_format.space_before = Pt(2)
    style.paragraph_format.space_after = Pt(2)
    style.paragraph_format.line_spacing = 1.15
    
    # Description style
    desc_style = doc.styles.add_style('Description', WD_STYLE_TYPE.PARAGRAPH)
    desc_style.font.name = 'Times New Roman'
    desc_style.font.size = Pt(12)
    desc_style.font.color.rgb = RGBColor(30, 30, 30)
    desc_style.paragraph_format.space_before = Pt(6)
    desc_style.paragraph_format.space_after = Pt(6)
    desc_style.paragraph_format.line_spacing = 1.5
    
    return style, desc_style


def add_code_block(doc, code_text, lang="TypeScript"):
    """Add formatted code block to document."""
    # File info line
    p = doc.add_paragraph()
    p.style = doc.styles['Description']
    run = p.add_run(f"Bahasa: {lang}")
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(100, 100, 100)
    
    # Code content
    lines = code_text.split('\n')
    for line in lines:
        p = doc.add_paragraph()
        p.style = doc.styles['CodeBlock']
        p.add_run(line)


def add_section(doc, section_data):
    """Add a complete section (heading + description + code) to the document."""
    if 'subsections' in section_data:
        # Section with subsections (e.g., IndoBERT)
        doc.add_heading(section_data['section'], level=1)
        
        for sub in section_data['subsections']:
            doc.add_heading(sub['subtitle'], level=2)
            
            # Description
            for para_text in sub['description'].split('\n\n'):
                p = doc.add_paragraph()
                p.style = doc.styles['Description']
                p.add_run(para_text)
            
            # File path info
            rel_path = os.path.relpath(sub['file'], BASE_DIR)
            p = doc.add_paragraph()
            p.style = doc.styles['Description']
            run = p.add_run(f"📄 File: {rel_path}")
            run.bold = True
            run.font.size = Pt(10)
            
            # Code
            if sub['file'].endswith('.ipynb'):
                code = extract_notebook_code(sub['file'])
            else:
                with open(sub['file'], 'r', encoding='utf-8') as f:
                    code = f.read()
            
            add_code_block(doc, code, sub['lang'])
            doc.add_page_break()
    else:
        # Simple section
        doc.add_heading(section_data['section'], level=1)
        
        # Description
        for para_text in section_data['description'].split('\n\n'):
            p = doc.add_paragraph()
            p.style = doc.styles['Description']
            p.add_run(para_text)
        
        # File path info
        rel_path = os.path.relpath(section_data['file'], BASE_DIR)
        p = doc.add_paragraph()
        p.style = doc.styles['Description']
        run = p.add_run(f"📄 File: {rel_path}")
        run.bold = True
        run.font.size = Pt(10)
        
        # Code
        with open(section_data['file'], 'r', encoding='utf-8') as f:
            code = f.read()
        
        add_code_block(doc, code, section_data['lang'])
        doc.add_page_break()


def main():
    doc = Document()
    
    # ─── Page setup ───
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)
    
    # ─── Create custom styles ───
    create_code_style(doc)
    
    # ─── Title Page ───
    doc.add_paragraph()  # spacing
    doc.add_paragraph()
    doc.add_paragraph()
    
    title = doc.add_heading('', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('Dokumentasi Kode Program')
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor(0, 51, 102)
    
    subtitle = doc.add_heading('', level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Pembobotan Fuzzy AHP, Rule-Based Engine,\ndan IndoBERT Intent Classifier')
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(51, 102, 153)
    
    doc.add_paragraph()
    
    # Project info
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.style = doc.styles['Description']
    run = info.add_run('Sistem Penilaian Risiko Halal Daging Sapi\n')
    run.font.size = Pt(14)
    run = info.add_run('(Knowledge Management System — Decision Support System)')
    run.font.size = Pt(12)
    run.italic = True
    
    doc.add_paragraph()
    
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.style = doc.styles['Description']
    run = date_para.add_run(f"Dibuat: {datetime.now().strftime('%d %B %Y, %H:%M')}")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(100, 100, 100)
    
    doc.add_page_break()
    
    # ─── Table of Contents (Manual) ───
    doc.add_heading('Daftar Isi', level=1)
    
    toc_items = [
        "1. Pembobotan Fuzzy AHP (fuzzyAHP.ts)",
        "   • Tipe TFN dan Skala Fuzzy Linguistik",
        "   • Fungsi Matematika Inti (FSE, Defuzzifikasi, Normalisasi)",
        "   • Consistency Ratio (CR)",
        "   • Integrasi Database (Prisma ORM)",
        "",
        "2. Rule-Based Risk Assessment Engine (rule-engine.ts)",
        "   • Tipe Data dan Interface",
        "   • Loader Rule Base (Singleton)",
        "   • Evaluasi Konstruk, Stage, dan Overall",
        "   • Agregasi Weakest-Link (MAX)",
        "",
        "3. IndoBERT Intent Classifier",
        "   3.1. Inference Runtime (intent-classifier.ts)",
        "       • Singleton Pattern untuk Model Loading",
        "       • Klasifikasi Intent",
        "   3.2. Training / Fine-tuning (train_indobert_onnx_v2.ipynb)",
        "       • Dataset Preparation",
        "       • Tokenisasi dan Model Loading",
        "       • Training Arguments dan Fine-tuning",
        "       • Ekspor ONNX dan Upload ke Hugging Face",
        "",
        "4. Unit Test — Fuzzy AHP (fuzzyAHP.test.ts)",
    ]
    
    for item in toc_items:
        p = doc.add_paragraph()
        p.style = doc.styles['Description']
        if item.startswith("   "):
            p.paragraph_format.left_indent = Cm(1.5)
            if item.strip().startswith("•"):
                p.paragraph_format.left_indent = Cm(2.5)
        p.add_run(item.strip())
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
    
    doc.add_page_break()
    
    # ─── Content Sections ───
    for section_data in FILES:
        add_section(doc, section_data)
    
    # ─── Save ───
    doc.save(OUTPUT_PATH)
    print(f"[DONE] DOCX berhasil dibuat: {OUTPUT_PATH}")
    return OUTPUT_PATH


if __name__ == '__main__':
    main()
